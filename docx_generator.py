"""
Professional DOCX kurs ishi generator
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import io


def _set_paragraph_spacing(paragraph, before=0, after=6, line_spacing=1.5):
    """Paragraf oraliqlarini sozlash"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def _add_heading(doc, text: str, level: int = 1, font_size: Pt = Pt(16),
                 color: RGBColor = None, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Sarlavha qo'shish"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = alignment
    for run in heading.runs:
        run.font.size = font_size
        run.font.name = 'Times New Roman'
        if color:
            run.font.color.rgb = color
    _set_paragraph_spacing(heading, before=12, after=8, line_spacing=1.5)
    return heading


def _add_paragraph(doc, text: str, font_size: Pt = Pt(14), bold: bool = False,
                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent: Cm = Cm(1.25),
                   color: RGBColor = None):
    """Oddiy paragraf qo'shish"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.first_line_indent = first_line_indent
    _set_paragraph_spacing(p, before=0, after=6, line_spacing=1.5)

    run = p.add_run(text)
    run.font.size = font_size
    run.font.name = 'Times New Roman'
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

    return p


def generate_professional_docx(course_data: dict) -> bytes:
    """
    Professional DOCX kurs ishi generatsiya qilish

    course_data = {
        "title": "Kurs ishi nomi",
        "subject": "Fan nomi",
        "introduction": "Kirish...",
        "chapters": [
            {"title": "1-bob...", "paragraphs": ["...", "..."]},
        ],
        "conclusion": "Xulosa...",
        "references": ["...", "..."]
    }
    """
    doc = Document()

    # Sahifa sozlamalari
    section = doc.sections[0]
    section.page_width = Cm(21.0)  # A4
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)

    # Default shrift
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5

    # ===== TITUL SAHIFASI =====
    # Bo'sh joy
    for _ in range(3):
        doc.add_paragraph()

    # Universitet nomi
    _add_paragraph(
        doc,
        "O'ZBEKISTON RESPUBLIKASI\nOLIY TA'LIM, FAN VA INNOVATSIYALAR VAZIRLIGI",
        font_size=Pt(14), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=Cm(0)
    )

    doc.add_paragraph()

    # Kurs ishi sarlavhasi
    _add_paragraph(
        doc, "KURS ISHI",
        font_size=Pt(18), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=Cm(0)
    )

    doc.add_paragraph()

    # Mavzu
    _add_paragraph(
        doc, f"Mavzu: {course_data.get('title', '')}",
        font_size=Pt(16), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=Cm(0)
    )

    doc.add_paragraph()
    _add_paragraph(
        doc, f"Fan: {course_data.get('subject', '')}",
        font_size=Pt(14), alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=Cm(0)
    )

    for _ in range(4):
        doc.add_paragraph()

    # Bajardi / Qabul qildi
    _add_paragraph(
        doc, "Bajardi: _________________________",
        font_size=Pt(14), alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent=Cm(0)
    )
    _add_paragraph(
        doc, "Qabul qildi: ______________________",
        font_size=Pt(14), alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        first_line_indent=Cm(0)
    )

    doc.add_paragraph()
    _add_paragraph(
        doc, "Toshkent — 2025",
        font_size=Pt(14), alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent=Cm(0)
    )

    # Yangi sahifa
    doc.add_page_break()

    # ===== MUNDARIJA =====
    _add_heading(
        doc, "MUNDARIJA", level=1, font_size=Pt(16),
        alignment=WD_ALIGN_PARAGRAPH.CENTER
    )
    doc.add_paragraph()

    toc_items = [
        ("KIRISH", "3"),
        ("1-BOB. NAZARIY ASOSLAR", "4"),
    ]
    if course_data.get('chapters'):
        for i, ch in enumerate(course_data['chapters']):
            title_clean = ch.get('title', f'{i+1}-bob')
            page = str(4 + i * 3)
            toc_items.append((title_clean, page))

    toc_items += [
        ("XULOSA", str(4 + len(course_data.get('chapters', [])) * 3)),
        ("FOYDALANILGAN ADABIYOTLAR", str(6 + len(course_data.get('chapters', [])) * 3)),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_paragraph_spacing(p, before=2, after=2, line_spacing=1.5)
        run = p.add_run(f"{item} {'.' * (50 - len(item))} {page}")
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'

    # ===== KIRISH =====
    doc.add_page_break()
    _add_heading(
        doc, "KIRISH", level=1, font_size=Pt(16),
        alignment=WD_ALIGN_PARAGRAPH.CENTER
    )

    intro = course_data.get('introduction', '')
    if intro:
        for paragraph_text in intro.split('\n'):
            if paragraph_text.strip():
                _add_paragraph(doc, paragraph_text.strip())

    # Dolzarblik, maqsad, vazifalar
    _add_paragraph(
        doc, f"Mazkur kurs ishining dolzarbligi shundaki, {course_data.get('title', 'ushbu mavzu')} "
             f"bugungi kunda dolzarb ahamiyat kasb etmoqda.",
        font_size=Pt(14)
    )
    _add_paragraph(
        doc, "Kurs ishining maqsadi — ushbu mavzuni chuqur o'rganish va tahlil qilish asosida "
             "nazariy hamda amaliy xulosalar chiqarishdan iborat.",
        font_size=Pt(14)
    )
    _add_paragraph(
        doc, "Kurs ishining vazifalari:\n"
             "- Mavzuga oid adabiyotlarni o'rganish va tahlil qilish;\n"
             "- Nazariy asoslarni yoritish;\n"
             "- Amaliy tahlil o'tkazish;\n"
             "- Taklif va tavsiyalar ishlab chiqish.",
        font_size=Pt(14)
    )

    # ===== BOBLAR =====
    chapters = course_data.get('chapters', [])
    for ch_idx, chapter in enumerate(chapters):
        doc.add_page_break()

        ch_title = chapter.get('title', f'{ch_idx+1}-bob')
        _add_heading(
            doc, ch_title.upper(), level=1, font_size=Pt(16),
            alignment=WD_ALIGN_PARAGRAPH.CENTER
        )
        doc.add_paragraph()

        paragraphs = chapter.get('paragraphs', [])
        for p_idx, para_text in enumerate(paragraphs):
            if para_text.strip():
                # Har bir bo'lim sarlavhasi
                if p_idx == 0:
                    _add_paragraph(
                        doc, f"{ch_idx+1}.{p_idx+1}. {para_text[:80]}{'...' if len(para_text) > 80 else ''}",
                        font_size=Pt(14), bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        first_line_indent=Cm(1.25)
                    )
                    doc.add_paragraph()

                _add_paragraph(doc, para_text.strip())

                # Qo'shimcha izoh va tahlil
                if len(para_text) > 200 and p_idx < len(paragraphs) - 1:
                    _add_paragraph(
                        doc,
                        f"Yuqoridagi fikrlardan kelib chiqib aytish mumkinki, ushbu yo'nalishdagi "
                        f"tadqiqotlar natijalari {course_data.get('title', 'mavzu')} bo'yicha "
                        f"muhim ilmiy-amaliy ahamiyatga ega.",
                        font_size=Pt(14)
                    )

    # ===== XULOSA =====
    doc.add_page_break()
    _add_heading(
        doc, "XULOSA", level=1, font_size=Pt(16),
        alignment=WD_ALIGN_PARAGRAPH.CENTER
    )

    conclusion = course_data.get('conclusion', '')
    if conclusion:
        for para_text in conclusion.split('\n'):
            if para_text.strip():
                _add_paragraph(doc, para_text.strip())

    _add_paragraph(
        doc,
        f"Xulosa qilib aytganda, {course_data.get('title', 'ushbu mavzu')} bo'yicha olib borilgan "
        f"tadqiqot natijalari shuni ko'rsatadiki, mazkur yo'nalishda yanada chuqurroq ilmiy "
        f"izlanishlar olib borish maqsadga muvofiqdir. "
        f"Kurs ishi davomida olingan natijalar amaliyotda qo'llanilishi mumkin.",
        font_size=Pt(14)
    )

    # ===== ADABIYOTLAR =====
    doc.add_page_break()
    _add_heading(
        doc, "FOYDALANILGAN ADABIYOTLAR", level=1, font_size=Pt(16),
        alignment=WD_ALIGN_PARAGRAPH.CENTER
    )
    doc.add_paragraph()

    references = course_data.get('references', [])
    if not references:
        references = [
            "O'zbekiston Respublikasi Konstitutsiyasi. — T.: O'zbekiston, 2023.",
            "Mirziyoyev Sh.M. Yangi O'zbekiston taraqqiyot strategiyasi. — T.: O'zbekiston, 2022.",
            "Karimov I.A. Yuksak ma'naviyat — yengilmas kuch. — T.: Ma'naviyat, 2008.",
            "Axmedov B. va boshq. Zamonaviy tadqiqot metodologiyasi. — T.: Fan, 2021.",
            "Raximov S. Ilmiy izlanish asoslari. — T.: Universitet, 2020.",
            "Smith J. Research Methods in Practice. — London: Academic Press, 2019.",
            "Johnson M. Data Analysis Techniques. — New York: Springer, 2020.",
            "Brown A. Modern Scientific Approaches. — Berlin: De Gruyter, 2021.",
        ]

    for i, ref in enumerate(references, 1):
        _add_paragraph(
            doc, f"{i}. {ref}",
            font_size=Pt(14), first_line_indent=Cm(0),
            alignment=WD_ALIGN_PARAGRAPH.LEFT
        )

    # Saqlash
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()
